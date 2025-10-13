"""
bilingual_fill.py
Make Word docs fully bilingual (VI → EN or EN → VI) while preserving look.
Rules:
- Order: Vietnamese first, then English
- English line: italic + Accent 1 (blue)
- Works even if everyone uses "Normal" style (infers base look from dominant run)
- Preserves bold/italic spans exactly (via <b>/<i>/<bi> tags)
- Handles paragraphs + tables
"""

import re
from copy import deepcopy
from typing import Optional, Tuple, List

from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from langdetect import detect, DetectorFactory

# Optional: plug in your translator. Default uses a no-op stub; Streamlit wraps with OpenAI.
# Provide a function translate_vi_en(text, src_lang)-> str  where src_lang in {"vi","en"}.
# You can swap this with any engine (OpenAI, Azure, DeepL, etc.)
def default_translate(text: str, src_lang: str) -> str:
    # Placeholder – you MUST inject a real translator from your app.
    return text

DetectorFactory.seed = 0

# -------------------
# Language utils
# -------------------
def detect_lang(text: str) -> Optional[str]:
    s = (text or "").strip()
    if not s:
        return None
    try:
        lang = detect(s)
        return lang if lang in {"vi", "en"} else None
    except Exception:
        return None

def needs_translation(text: str) -> Tuple[bool, Optional[str]]:
    s = (text or "").strip()
    if not s:
        return (False, None)
    lang = detect_lang(s)
    if lang not in {"vi", "en"}:
        return (False, None)
    return (True, lang)

# -------------------
# Emphasis tagging (bold/italic)
# -------------------
def paragraph_spans(par: Paragraph) -> List[Tuple[str, bool, bool]]:
    spans = []
    cur_text, cur_b, cur_i = "", None, None
    for r in par.runs:
        t = r.text or ""
        b = bool(r.bold)
        i = bool(r.italic)
        if cur_text and (b == cur_b and i == cur_i):
            cur_text += t
        else:
            if cur_text:
                spans.append((cur_text, cur_b, cur_i))
            cur_text, cur_b, cur_i = t, b, i
    if cur_text:
        spans.append((cur_text, cur_b, cur_i))
    return spans

def cell_spans(cell: _Cell) -> List[Tuple[str, bool, bool]]:
    spans = []
    cur_text, cur_b, cur_i = "", None, None
    for p in cell.paragraphs:
        for r in p.runs:
            t = r.text or ""
            b = bool(r.bold)
            i = bool(r.italic)
            if cur_text and (b == cur_b and i == cur_i):
                cur_text += t
            else:
                if cur_text:
                    spans.append((cur_text, cur_b, cur_i))
                cur_text, cur_b, cur_i = t, b, i
        if cur_text:
            spans.append((cur_text, cur_b, cur_i)); cur_text = ""
    return spans

def tag_spans(spans: List[Tuple[str, bool, bool]]) -> str:
    pieces = []
    for txt, b, i in spans:
        if not txt:
            continue
        if b and i:
            pieces.append(f"<bi>{txt}</bi>")
        elif b:
            pieces.append(f"<b>{txt}</b>")
        elif i:
            pieces.append(f"<i>{txt}</i>")
        else:
            pieces.append(txt)
    return "".join(pieces)

TAG_TOKEN = re.compile(r"(</?bi>|</?b>|</?i>)")

def parse_tagged_to_segments(tagged_text: str) -> List[Tuple[str, bool, bool]]:
    bold = italic = False
    segs: List[Tuple[str, bool, bool]] = []
    buf: List[str] = []

    def flush():
        if buf:
            segs.append(("".join(buf), bold, italic))
            buf.clear()

    tokens = TAG_TOKEN.split(tagged_text)
    for tk in tokens:
        if not tk:
            continue
        if tk == "<b>":
            flush(); bold = True
        elif tk == "</b>":
            flush(); bold = False
        elif tk == "<i>":
            flush(); italic = True
        elif tk == "</i>":
            flush(); italic = False
        elif tk == "<bi>":
            flush(); bold = True; italic = True
        elif tk == "</bi>":
            flush(); bold = False; italic = False
        else:
            buf.append(tk)
    flush()
    return segs

# -------------------
# Visual inference & cloning
# -------------------
def font_size_points(run: Run) -> Optional[float]:
    try:
        return float(run.font.size.pt) if (run.font.size and hasattr(run.font.size, "pt")) else None
    except Exception:
        return None

def visual_weight(run: Run) -> float:
    text = (run.text or "").strip()
    if not text:
        return 0.0
    score = 0.0
    size = font_size_points(run)
    if size:
        score += size
    if run.bold:
        score += 3.0
    if run.underline:
        score += 1.5
    if run.italic:
        score += 1.0
    letters = [c for c in text if c.isalpha()]
    if letters:
        caps_ratio = sum(1 for c in letters if c.isupper()) / len(letters)
        if caps_ratio > 0.7 and len(letters) >= 4:
            score += 2.5
    score += min(len(text) / 40.0, 2.0)
    return score

def pick_donor_run_from_paragraph(par: Paragraph) -> Optional[Run]:
    if not par.runs:
        return None
    return max(par.runs, key=visual_weight)

def pick_donor_run_from_cell(cell: _Cell) -> Optional[Run]:
    cand = None
    best = -1.0
    for p in cell.paragraphs:
        for r in p.runs:
            w = visual_weight(r)
            if w > best:
                best, cand = w, r
    return cand

def clone_run_format(src_run: Run, dst_run: Run):
    """
    Copy run rPr (bold/italic/size/font/color/etc.) by cloning XML.
    """
    src_r = src_run._r
    dst_r = dst_run._r
    src_rPr = src_r.rPr
    if src_rPr is not None:
        dst_rPr = deepcopy(src_rPr)
        if dst_r.rPr is not None:
            dst_r.remove(dst_r.rPr)
        dst_r.append(dst_rPr)

def copy_paragraph_format(src: Paragraph, dst: Paragraph):
    try:
        pf_src = src.paragraph_format
        pf_dst = dst.paragraph_format
        pf_dst.left_indent = pf_src.left_indent
        pf_dst.right_indent = pf_src.right_indent
        pf_dst.first_line_indent = pf_src.first_line_indent
        pf_dst.space_before = pf_src.space_before
        pf_dst.space_after = pf_src.space_after
        pf_dst.line_spacing = pf_src.line_spacing
        pf_dst.alignment = pf_src.alignment
    except Exception:
        pass

# -------------------
# Guards & prefixes
# -------------------
def paragraph_already_followed_by_translation(par: Paragraph) -> bool:
    parent = par._p.getparent()
    idx = parent.index(par._p)
    if idx + 1 < len(parent):
        nxt = Paragraph(parent[idx + 1], par._parent)
        a = (par.text or "").strip().lower()
        b = (nxt.text or "").strip().lower().replace("[en] ", "").replace("[vi] ", "")
        return a == b
    return False

LIST_PREFIX_RE = re.compile(r"""
    ^\s*(
        (\(?[0-9ivxlcdmIVXLCDM]+\)?[.)]?) |   # 1. 1) (1) I. i)
        ([A-Za-z]\.) |                        # A. a.
        ([-•*])                               # -, •, *
    )\s+
""", re.VERBOSE)

def extract_list_prefix(text: str) -> Tuple[str, str]:
    m = LIST_PREFIX_RE.match(text or "")
    if m:
        prefix = m.group(0)
        rest = (text or "")[len(prefix):]
        return prefix, rest
    return "", text or ""

# -------------------
# Theme helpers for EN line
# -------------------
def apply_english_style(run: Run):
    """Make English run italic + Accent 1 (blue)."""
    run.italic = True
    try:
        run.font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_1
    except Exception:
        # Fallback: plain blue-ish RGB if theme missing
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(31, 73, 125)  # similar to Office Accent 1

# -------------------
# Writers
# -------------------
def clear_paragraph(par: Paragraph):
    # remove all runs
    for r in list(par.runs):
        r._r.getparent().remove(r._r)

def write_runs_from_segments(par: Paragraph, segments: List[Tuple[str, bool, bool]],
                             donor: Optional[Run], override_italic: Optional[bool] = None,
                             color_accent1_for_all: bool = False):
    """Write multiple runs into 'par' from segments, cloning base from donor.
       If override_italic is not None, force italic per run.
       If color_accent1_for_all, apply Accent 1 to all runs.
    """
    for txt, b, i in segments:
        if not txt:
            continue
        r = par.add_run()
        if donor:
            clone_run_format(donor, r)
        # emphasis from source unless overridden
        r.bold = True if b else False
        r.italic = (override_italic if override_italic is not None else (True if i else False))
        if color_accent1_for_all:
            apply_english_style(r)  # sets italic True; we might need to re-apply italic override
            if override_italic is not None:
                r.italic = override_italic
        r.text = txt

def insert_new_par_after(src_par: Paragraph) -> Paragraph:
    p = src_par._p
    new_p_xml = OxmlElement("w:p")
    p.addnext(new_p_xml)
    new_par = Paragraph(new_p_xml, src_par._parent)
    copy_paragraph_format(src_par, new_par)
    if src_par.style is not None:
        new_par.style = src_par.style
    return new_par

# Core rule:
# - If source VI: keep current paragraph as VI; insert EN below (italic + accent1)
# - If source EN: replace current paragraph with VI; insert EN below (from original EN) italic + accent1
def make_paragraph_bilingual(
    par: Paragraph,
    translate_fn = default_translate,
    add_labels: bool = True
):
    src_text = par.text or ""
    need, lang = needs_translation(src_text)
    if not need:
        return
    # Build tagged segments from SOURCE paragraph (for accurate emphasis)
    src_spans = paragraph_spans(par)
    tagged_src = tag_spans(src_spans)
    donor = pick_donor_run_from_paragraph(par) or (par.runs[0] if par.runs else None)

    if lang == "vi":
        # VI is already first; we only add EN line below.
        if paragraph_already_followed_by_translation(par):
            return
        # Translate VI -> EN (preserve tags)
        translated_tagged = translate_fn(tagged_src, "vi")
        segs_en = parse_tagged_to_segments(translated_tagged)
        # Create EN paragraph below, with EN styling
        new_par = insert_new_par_after(par)
        if add_labels:
            new_par.add_run("[EN] ")
        write_runs_from_segments(new_par, segs_en, donor, override_italic=True, color_accent1_for_all=True)

    elif lang == "en":
        # Need VI first. We'll overwrite the current paragraph with VI,
        # then add the original EN below (italic + accent1).
        translated_tagged = translate_fn(tagged_src, "en")
        segs_vi = parse_tagged_to_segments(translated_tagged)
        # Overwrite current paragraph with VI
        vi_par = par
        if add_labels:
            clear_paragraph(vi_par)
            vi_par.add_run("[VI] ")
        else:
            clear_paragraph(vi_par)
        write_runs_from_segments(vi_par, segs_vi, donor, override_italic=False, color_accent1_for_all=False)

        # EN line below using ORIGINAL source segments (tagged_src)
        segs_en_src = parse_tagged_to_segments(tagged_src)  # original EN emphasis
        en_par = insert_new_par_after(vi_par)
        if add_labels:
            en_par.add_run("[EN] ")
        write_runs_from_segments(en_par, segs_en_src, donor, override_italic=True, color_accent1_for_all=True)

def make_cell_bilingual(
    cell: _Cell,
    translate_fn = default_translate,
    add_labels: bool = True,
    table_mode: str = "stack"
):
    raw = (cell.text or "").strip()
    if not raw:
        return
    need, lang = needs_translation(raw)
    if not need:
        return

    donor = pick_donor_run_from_cell(cell)
    spans_src = cell_spans(cell)
    tagged_src = tag_spans(spans_src)

    # Clear cell and rebuild single paragraph
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    par = cell.add_paragraph()

    if lang == "vi":
        # VI first (original), then EN (italic blue)
        # Write VI (from source spans)
        if add_labels:
            par.add_run("[VI] ")
        write_runs_from_segments(par, parse_tagged_to_segments(tagged_src), donor, override_italic=False)

        if table_mode == "stack":
            par.add_run("\n")
        else:
            par.add_run(" / ")

        # EN translated
        segs_en = parse_tagged_to_segments(translate_fn(tagged_src, "vi"))
        if add_labels:
            par.add_run("[EN] ")
        write_runs_from_segments(par, segs_en, donor, override_italic=True, color_accent1_for_all=True)

    elif lang == "en":
        # Need VI first. Translate EN->VI, then write EN (original) second with italic blue
        segs_vi = parse_tagged_to_segments(translate_fn(tagged_src, "en"))
        if add_labels:
            par.add_run("[VI] ")
        write_runs_from_segments(par, segs_vi, donor, override_italic=False)

        if table_mode == "stack":
            par.add_run("\n")
        else:
            par.add_run(" / ")

        if add_labels:
            par.add_run("[EN] ")
        write_runs_from_segments(par, parse_tagged_to_segments(tagged_src), donor, override_italic=True, color_accent1_for_all=True)

def process_docx(
    in_path: str,
    out_path: str,
    translate_fn = default_translate,
    add_labels: bool = True,
    table_mode: str = "stack"
):
    doc = Document(in_path)

    # paragraphs
    for par in list(doc.paragraphs):
        make_paragraph_bilingual(par, translate_fn=translate_fn, add_labels=add_labels)

    # tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                make_cell_bilingual(cell, translate_fn=translate_fn, add_labels=add_labels, table_mode=table_mode)

    doc.save(out_path)
